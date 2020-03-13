# qar5_docx.py (qar5)
# !/usr/bin/env python3
# coding=utf-8
# young.daniel@epa.gov

"""Definition of qar5 docx export methods."""

from docx import Document
from docx.enum.style import WD_BUILTIN_STYLE, WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX, WD_LINE_SPACING, \
    WD_PARAGRAPH_ALIGNMENT
    
from docx.shared import Inches
from os import path
# import pypandoc
# import tempfile
# from zipfile import ZipFile
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, HttpResponse
from django.template.loader import render_to_string
from django.templatetags.static import static
from DataSearch.settings import DEBUG, STATIC_ROOT
from qar5.views import get_all_qar5_for_user, get_qapp_info

def export_doc(request, *args, **kwargs):
    """Function to export QAR5 as a Word Document."""
    qapp_id = kwargs.get('pk', None) 
    template_name = 'export/qar5_pdf_template.html'

    if qapp_id is None:
        # TODO: Export ALL QAR5 objects available for this user to Docx.
        qapp_info = get_all_qar5_for_user(request.user)

    else:
        qapp_info = get_qapp_info(request.user, qapp_id)
        if not qapp_info:
            return

        filename = '%s.docx' % qapp_info['qapp'].title

        document = Document()
        styles = document.styles
        
        # #################################################
        # BEGIN COVER PAGE
        # #################################################
        # Coversheet with signatures section:
        if DEBUG:
            logo = path.join(STATIC_ROOT, 'EPA_Files', 'logo.png')
        else:
            logo = static('logo.png')

        document.add_picture(logo, width=Inches(1.25))

        # row has right-aligned box "Quality Assurance Project Plan"
        # Align the picture/text WD_ALIGN_PARAGRAPH.RIGHT
        blue_header_style = document.styles.add_style(
            'blue_header', WD_STYLE_TYPE.PARAGRAPH).paragraph_format
        blue_header_style.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        blue_header = document.add_heading('Quality Assurance Project Plan',
                                           level=1)

        style_center = document.styles.add_style(
            'center_text', WD_STYLE_TYPE.PARAGRAPH).paragraph_format
        style_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        style_paragraph = document.styles.add_style(
            'style_paragraph', WD_STYLE_TYPE.PARAGRAPH).paragraph_format
        style_paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # TODO Make blue_header text white, add blue background with shadow
        # document.add_picture('blue_background.png', width=Inches(4))
        # background color: rgb(0, 176, 240)

        # The rest of the document will be WD_ALIGN_PARAGRAPH.CENTER

        # blank line
        document.add_heading('Office of Research and Development', level=1)
        document.add_heading(qapp_info['qapp'].division.name, level=1)
        # blank line
        # Next few sections are from the qapp object
        document.add_heading(qapp_info['qapp'].division_branch, level=3)
        # blank line
        document.add_heading('EPA Project Lead', level=2)
        for lead in qapp_info['qapp_leads']:
            document.add_heading(lead.name, level=3)
        # blank line
        document.add_heading(qapp_info['qapp'].intra_extra, level=3)
        document.add_heading(qapp_info['qapp'].qa_category, level=3)
        document.add_heading(qapp_info['qapp'].revision_number, level=3)
        document.add_heading(str(qapp_info['qapp'].date), level=3)
        # blank line
        document.add_heading('Prepared By', level=2)
        document.add_heading(
            '%s %s' % (qapp_info['qapp'].prepared_by.first_name,
                       qapp_info['qapp'].prepared_by.last_name,),
            level=3)
        # blank line
        document.add_heading(qapp_info['qapp'].strap, level=3)
        document.add_heading(qapp_info['qapp'].tracking_id, level=3)

        # #################################################
        # END COVER PAGE
        # BEGIN APPROVAL PAGE
        # #################################################

        document.add_heading('Approval Page', level=2)
        # Signature grid ...
        num_signatures = len(qapp_info['signatures'])
        table = document.add_table(rows=6+num_signatures, cols=12)
        table.style = styles['Table Grid']

        row_cells = table.rows[0].cells
        row_cells[0].text = 'QA Project Plan Title:'
        row_cells[0].merge(row_cells[3])
        row_cells[4].text = qapp_info['qapp_approval'].project_plan_title
        row_cells[4].merge(row_cells[11])

        row_cells = table.rows[1].cells
        row_cells[0].text = 'QA Activity Number:'
        row_cells[0].merge(row_cells[3])
        row_cells[4].text = qapp_info['qapp_approval'].activity_number
        row_cells[4].merge(row_cells[11])

        # TODO: Center text in this row:
        row_cells = table.rows[2].cells
        row_cells[0].text = 'If Intramural or Extramural, EPA Project Approvals'
        row_cells[0].merge(row_cells[11])

        iter_count = 0
        # TODO: Iterate through EPA Project Approvals:
        # Start with row 3 + iter_count++
        for sig in qapp_info['signatures']:
            if not sig.contractor:
                row_cells = table.rows[3 + iter_count].cells
                row_cells[0].text = 'Name:'
                row_cells[1].text = sig.name
                row_cells[1].merge(row_cells[3])

                row_cells[4].text = 'Signature/Date:'
                row_cells[4].merge(row_cells[5])
                row_cells[6].merge(row_cells[11])
                iter_count += 1

        # Always insert a blank entry for hand-written approval sigs
        row_cells = table.rows[3 + iter_count].cells
        row_cells[0].text = 'Name:'
        row_cells[1].merge(row_cells[3])
        row_cells[4].text = 'Signature/Date:'
        row_cells[4].merge(row_cells[5])
        row_cells[6].merge(row_cells[11])

        # TODO: Center text in this row:
        row_cells = table.rows[4 + iter_count].cells
        row_cells[0].text = 'If Extramural, Contractor Project Approvals'
        row_cells[0].merge(row_cells[11])

        # TODO: Iterate through Contractor Project Approvals:
        # Start with row 5 + iter_count++
        for sig in qapp_info['signatures']:
            if sig.contractor:
                row_cells = table.rows[5 + iter_count].cells
                row_cells[0].text = 'Name:'
                row_cells[1].text = sig.name
                row_cells[1].merge(row_cells[3])

                row_cells[4].text = 'Signature/Date:'
                row_cells[4].merge(row_cells[5])
                row_cells[6].merge(row_cells[11])
                iter_count += 1

        # Always insert a blank entry for hand-written approval sigs
        row_cells = table.rows[5 + iter_count].cells
        row_cells[0].text = 'Name:'
        row_cells[1].merge(row_cells[3])
        row_cells[4].text = 'Signature/Date:'
        row_cells[4].merge(row_cells[5])
        row_cells[6].merge(row_cells[11])
        document.add_page_break()

        # #################################################
        # END APPROVAL PAGE
        # BEGIN ToC PAGE
        # #################################################

        document.add_heading('Table of Contents', level=2)
        document.add_heading(
            'TODO: This will have to be generated after the rest of ' + \
            'the doc so we know page numbers and contents...', level=3)
        document.add_page_break()

        # #################################################
        # END ToC PAGE
        # BEGIN Everything Else PAGE
        # #################################################

        #  1) Heading 1 - Revision History
        document.add_heading('Revision History', level=1)
        #  2) Table Label
        document.add_heading('Revision History', level=3)
        #  3) Table (revision history)
        # TODO

        # TODO: Paragraphs aren't formatting properly, still double spaces...
        document.add_heading('Section A - Executive Summary', level=1)
        document.add_heading('A.1 Distribution List', level=2)
        document.add_paragraph(qapp_info['section_a'].a3, 'No Spacing')
        document.add_heading('A.2 Project Task Organization', level=2)
        document.add_paragraph(qapp_info['section_a'].a4, 'No Spacing')
        document.add_heading('A.3 Problem Definition Background', level=2)
        document.add_paragraph(qapp_info['section_a'].a5, 'No Spacing')
        document.add_heading('A.4 Project Description', level=2)
        document.add_paragraph(qapp_info['section_a'].a6, 'No Spacing')
        document.add_heading('A.5 Quality Objectives and Criteria', level=2)
        document.add_paragraph(qapp_info['section_a'].a7, 'No Spacing')
        document.add_heading('A.6 Special Training Certification', level=2)
        document.add_paragraph(qapp_info['section_a'].a8, 'No Spacing')
        document.add_heading('A.7 Documents and Records', level=2)
        document.add_paragraph(qapp_info['section_a'].a9, 'No Spacing')

        #  1) Heading 1 - Section B - Experimental Design
        document.add_heading('Section B - Experimental Design', level=1)
        #  2) Heading 2 - B.1 Sample/Data Collection, Gathering, or Use
        document.add_heading('B.1 Sample/Data Collection, Gathering, or Use',
                             level=2)
        #  3) Heading 3 - B.1.1 Use
        document.add_heading('B.1.1 Use', level=3)
        #  4) Paragraph - asdf
        document.add_paragraph(qapp_info['section_b'].b1_2, 'No Spacing')
        #  5) Heading 3 - B.1.2 Requirements
        document.add_heading('B.1.2 Requirements', level=3)
        #  6) Paragraph - adsf
        document.add_paragraph(qapp_info['section_b'].b1_3, 'No Spacing')
        #  7) Heading 3 - B.1.3 Databases, Maps, Literature
        document.add_heading('B.1.3 Databases, Maps, Literature', level=3)
        #  8) Paragraph - adsf
        document.add_paragraph(qapp_info['section_b'].b1_4, 'No Spacing')
        #  9) Heading 3 - B.1.4 Non-Quality Constraints
        document.add_heading('B.1.4 Non-Quality Constraints', level=3)
        # 10) Paragraph - adsf
        document.add_paragraph(qapp_info['section_b'].b1_5, 'No Spacing')
        # 11) Heading 2 - B.2 Data Analysis / Statistical Design / Data Management
        document.add_heading(
            'B.2 Data Analysis / Statistical Design / Data Management',
            level=2)
        # 12) Heading 3 - B.2.1 Sources
        document.add_heading('B.2.1 Sources', level=3)
        # 13) Paragraph - asdf
        document.add_paragraph(qapp_info['section_b'].b2_1, 'No Spacing')
        # 14) Heading 3 - B.2.2 Acceptance/Rejection Process
        document.add_heading('B.2.2 Acceptance/Rejection Process', level=3)
        # 15) Paragraph - adf
        document.add_paragraph(qapp_info['section_b'].b2_2, 'No Spacing')
        # 16) Heading 3 - B.2.3 Rationale for Selections
        document.add_heading('B.2.3 Rationale for Selections', level=3)
        # 17) Paragraph - asdf
        document.add_paragraph(qapp_info['section_b'].b2_3, 'No Spacing')
        # 18) Heading 3 - B.2.4 Procedures
        document.add_heading('B.2.4 Procedures', level=3)
        # 19) Paragraph - adsf
        document.add_paragraph(qapp_info['section_b'].b2_4, 'No Spacing')
        # 20) Heading 3 - B.2.5 Disclaimer
        document.add_heading('B.2.5 Disclaimer', level=3)
        # 21) Paragraph - asdf
        document.add_paragraph(qapp_info['section_b'].b2_5, 'No Spacing')
        # 22) Heading 2 - B.3 Data Management and Documentation
        document.add_heading('B.3 Data Management and Documentation', level=2)
        # 23) Paragraph - asdf
        document.add_paragraph(qapp_info['section_b'].b3, 'No Spacing')
        # 24) Heading 2 - B.4 Tracking
        document.add_heading('B.4 Tracking', level=2)
        # 25) Paragraph - asfd
        document.add_paragraph(qapp_info['section_b'].b4, 'No Spacing')

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.' + \
                'wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=%s' % filename
        document.save(response)
        #document.save('test_export.docx')
        return response

    return